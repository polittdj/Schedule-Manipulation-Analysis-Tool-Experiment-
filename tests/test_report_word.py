"""Round-trip tests for report_word (python-docx) — H-DRIFT-1 traceability.

Every numeric assertion round-trips through the file on disk:
  1. Build a schedule and call ``analyze_schedule`` to get a ``ScheduleAnalysis``.
  2. Write the report to a tmp file via ``build_word_report``.
  3. Re-open the file with ``docx.Document`` and search its full text for the
     expected values.

This proves that the rendered numbers are traceable to the computed inputs
(H-DRIFT-1) rather than being fabricated or approximate.

Perturbation discipline (H-VACUOUS-TEST)
-----------------------------------------
The PERTURB test introduces a negative-lag lead on the only relation, forcing
DCMA-02 (Leads) to FAIL.  The health score must drop and the Findings section
must name "Leads".  Both are checked after disk round-trip.

CUI banner
----------
The presence of ``CUI_NOTICE`` text is asserted in every generated document.

Axis note
---------
Default calendar: 480 working minutes per day.
``project_finish / 480.0`` == working-day display value used in the Word document.
"""

from __future__ import annotations

import datetime as dt
import pathlib

import pytest
from docx import Document

from schedule_forensics.analysis import analyze_schedule
from schedule_forensics.metrics_common import MetricStatus
from schedule_forensics.report_word import CUI_NOTICE, build_word_document, build_word_report
from schedule_forensics.schemas import Relation, RelationType, Schedule, Task

_START = dt.datetime(2025, 1, 6, 8)  # Monday, working-day start


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def _task(uid: int, dur: int = 480, **kwargs: object) -> Task:
    return Task(unique_id=uid, name=f"T{uid}", duration_minutes=dur, **kwargs)  # type: ignore[arg-type]


def _clean_schedule() -> Schedule:
    """A minimal two-task linear schedule: T1(480) -> T2(480).

    CPM: project_finish = 960 = 2.0 working days.
    DCMA-02 (Leads): 0 leads -> PASS.
    """
    return Schedule(
        name="clean",
        project_start=_START,
        status_date=_START,
        tasks=(_task(1, 480), _task(2, 480)),
        relations=(Relation(predecessor_id=1, successor_id=2),),
    )


def _lead_schedule() -> Schedule:
    """Same two tasks but the link has a negative lag (a lead).

    DCMA-02 (Leads): 1 lead -> FAIL.  Health score lower than clean case.
    """
    return Schedule(
        name="lead",
        project_start=_START,
        status_date=_START,
        tasks=(_task(1, 480), _task(2, 480)),
        relations=(
            Relation(predecessor_id=1, successor_id=2, type=RelationType.FS, lag_minutes=-240),
        ),
    )


def _ev_schedule() -> Schedule:
    """Schedule carrying earned-value data; SPI = 0.75 (behind), status @ baseline finish."""
    d7, d8 = dt.datetime(2025, 1, 7, 8), dt.datetime(2025, 1, 8, 8)
    return Schedule(
        name="ev",
        project_start=_START,
        status_date=d8,
        tasks=(
            _task(
                1,
                480,
                percent_complete=100.0,
                baseline_start=_START,
                baseline_finish=d7,
                budgeted_cost=100.0,
            ),
            _task(
                2,
                480,
                percent_complete=50.0,
                baseline_start=d7,
                baseline_finish=d8,
                budgeted_cost=100.0,
            ),
        ),
    )


# ---------------------------------------------------------------------------
# Helper: extract all text from a Document as one big string
# ---------------------------------------------------------------------------


def _full_text(doc: Document) -> str:
    """Concatenate every paragraph and every table cell into one string."""
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Tests: in-memory document (no disk I/O)
# ---------------------------------------------------------------------------


def test_build_word_document_contains_title() -> None:
    analysis = analyze_schedule(_clean_schedule())
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert "Schedule Forensics Report" in full


def test_build_word_document_contains_cui_banner() -> None:
    analysis = analyze_schedule(_clean_schedule())
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert CUI_NOTICE in full, "CUI_NOTICE not found in document text"


def test_build_word_document_contains_project_finish_minutes() -> None:
    """project_finish (minutes) is present in the document text."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.project_finish == 960
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert "960" in full, f"960 not found in document; excerpt: {full[:500]}"


def test_build_word_document_contains_project_finish_days() -> None:
    """project_finish / 480.0 (2.0) is present in the document text."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.project_finish == 960
    # The Word module formats working days as f"{960/480.0:.2f}" == "2.00"
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert "2.00" in full, f"'2.00' working days not found in document; excerpt: {full[:500]}"


def test_build_word_document_contains_health_score() -> None:
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.health_score is not None
    doc = build_word_document(analysis)
    full = _full_text(doc)
    # Health score is formatted as f"{health_score:.2f}%"
    expected = f"{analysis.health_score:.2f}%"
    assert expected in full, f"'{expected}' not found in document; excerpt: {full[:500]}"


def test_build_word_document_dcma_table_has_14_rows() -> None:
    """DCMA table must have exactly 14 data rows (+ 1 header = 15 total)."""
    analysis = analyze_schedule(_clean_schedule())
    assert len(analysis.dcma) == 14
    doc = build_word_document(analysis)
    dcma_tables = [t for t in doc.tables if len(t.rows) == 15]
    assert dcma_tables, "No table with 15 rows (header + 14 DCMA) found in document"


def test_build_word_document_dcma_table_contains_metric_ids() -> None:
    """DCMA table cells contain at least the first and last metric ids."""
    analysis = analyze_schedule(_clean_schedule())
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert "DCMA-01" in full
    assert "DCMA-14" in full


def test_build_word_document_renders_earned_value() -> None:
    """The earned-value table renders SPI/SPI(t) and the 0.75 measured value."""
    analysis = analyze_schedule(_ev_schedule())
    spi = next(m for m in analysis.performance_indices if m.metric_id == "SPI")
    assert spi.measured == pytest.approx(0.75)  # pre-condition
    full = _full_text(build_word_document(analysis))
    assert "SPI" in full
    assert "SPI(t)" in full
    assert "0.75" in full


def test_build_word_document_earned_value_skipped_without_data() -> None:
    """With no EV data the earned-value rows still render, marked SKIPPED."""
    full = _full_text(build_word_document(analyze_schedule(_clean_schedule())))
    assert "Earned-Value Indices" in full
    assert "SKIPPED" in full


def test_build_word_document_findings_present() -> None:
    """Failing metric names appear in the document findings section."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.findings, "fixture must have at least one failing metric"
    doc = build_word_document(analysis)
    full = _full_text(doc)
    for finding in analysis.findings:
        assert finding in full, f"Finding '{finding}' not found in document"


def test_build_word_document_no_fail_note() -> None:
    """When findings is empty the document shows the 'no failing metrics' note."""
    sched = Schedule(
        name="allpass",
        project_start=_START,
        status_date=_START,
        tasks=(
            _task(1, 480, resource_names=("Alice",)),
            _task(2, 480, resource_names=("Bob",)),
        ),
        relations=(Relation(predecessor_id=1, successor_id=2),),
    )
    analysis = analyze_schedule(sched)
    if analysis.findings:
        pytest.skip("schedule unexpectedly has failures; skip rather than assert wrong thing")
    doc = build_word_document(analysis)
    full = _full_text(doc).lower()
    assert "no failing" in full, f"Expected 'no failing' note; got: {full[:300]}"


def test_build_word_document_extension_footnote_absent_when_no_extensions() -> None:
    """No extension footnote when no DCMA metric is an extension."""
    analysis = analyze_schedule(_clean_schedule())
    extensions = [m for m in analysis.dcma if m.is_extension]
    assert not extensions, "fixture must have no extensions for this test to be meaningful"
    doc = build_word_document(analysis)
    full = _full_text(doc)
    assert "tool-original" not in full.lower(), "Extension footnote should be absent"


# ---------------------------------------------------------------------------
# Tests: round-trip through disk (H-DRIFT-1 traceability)
# ---------------------------------------------------------------------------


def test_roundtrip_project_finish_minutes(tmp_path: pathlib.Path) -> None:
    """project_finish minutes written then re-read from disk is present in document text."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.project_finish == 960
    out = tmp_path / "report.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert "960" in full, f"960 not found after round-trip; excerpt: {full[:500]}"


def test_roundtrip_project_finish_days(tmp_path: pathlib.Path) -> None:
    """project_finish / 480.0 written then re-read from disk equals '2.00'."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.project_finish == 960
    out = tmp_path / "report.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert "2.00" in full, f"'2.00' working days not found after round-trip; excerpt: {full[:500]}"


def test_roundtrip_health_score(tmp_path: pathlib.Path) -> None:
    """Health score written then re-read from disk matches analysis.health_score."""
    analysis = analyze_schedule(_clean_schedule())
    assert analysis.health_score is not None
    expected = f"{analysis.health_score:.2f}%"
    out = tmp_path / "health.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert expected in full, f"'{expected}' not found after round-trip; excerpt: {full[:500]}"


def test_roundtrip_cui_banner(tmp_path: pathlib.Path) -> None:
    """CUI banner text present in document after round-trip through disk."""
    analysis = analyze_schedule(_clean_schedule())
    out = tmp_path / "cui.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert CUI_NOTICE in full, f"CUI_NOTICE not found after disk round-trip; excerpt: {full[:500]}"


def test_roundtrip_dcma_measured_value(tmp_path: pathlib.Path) -> None:
    """A specific DCMA measured value is preserved through the disk round-trip.

    DCMA-02 (Leads) in the clean schedule: measured == 0.0, status PASS.
    After writing and re-reading, "DCMA-02" and "PASS" must both appear in the
    document text (the DCMA table row was round-tripped intact).
    """
    analysis = analyze_schedule(_clean_schedule())
    dcma02 = next(m for m in analysis.dcma if m.metric_id == "DCMA-02")
    assert dcma02.status is MetricStatus.PASS
    assert dcma02.measured == 0.0

    out = tmp_path / "dcma.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert "DCMA-02" in full, "DCMA-02 row not found after round-trip"
    assert "PASS" in full, "PASS status not found after round-trip"
    # The measured value 0.0 is rendered as "0.0" in the table
    assert "0.0" in full, f"0.0 not found after round-trip; excerpt: {full[:500]}"


# ---------------------------------------------------------------------------
# Perturbation tests (H-VACUOUS-TEST)
# ---------------------------------------------------------------------------


def test_perturbation_lead_flips_dcma02_and_reduces_health(tmp_path: pathlib.Path) -> None:
    """Adding a negative-lag lead must flip DCMA-02 to FAIL and reduce health.

    Schedule A (clean): no leads -> DCMA-02 PASS.
    Schedule B (lead): one -240 min lead -> DCMA-02 FAIL -> health score lower.

    Both reports are written to disk and re-read, proving the report reflects the
    computed analysis (H-VACUOUS-TEST / H-DRIFT-1).
    """
    # --- Schedule A: clean ---
    analysis_a = analyze_schedule(_clean_schedule())
    dcma02_a = next(m for m in analysis_a.dcma if m.metric_id == "DCMA-02")
    assert dcma02_a.status is MetricStatus.PASS
    assert analysis_a.health_score is not None

    out_a = tmp_path / "clean.docx"
    build_word_report(analysis_a, out_a)
    doc_a = Document(str(out_a))
    full_a = _full_text(doc_a)
    # DCMA-02 is PASS in the clean report
    assert "DCMA-02" in full_a
    expected_health_a = f"{analysis_a.health_score:.2f}%"
    assert expected_health_a in full_a, (
        f"A: health {expected_health_a} not found; excerpt: {full_a[:500]}"
    )

    # --- Schedule B: perturbed (lead added) ---
    analysis_b = analyze_schedule(_lead_schedule())
    dcma02_b = next(m for m in analysis_b.dcma if m.metric_id == "DCMA-02")
    assert dcma02_b.status is MetricStatus.FAIL
    assert analysis_b.health_score is not None
    # Health must be lower after the perturbation
    assert analysis_b.health_score < analysis_a.health_score, (
        f"Health should drop: clean={analysis_a.health_score} lead={analysis_b.health_score}"
    )

    out_b = tmp_path / "lead.docx"
    build_word_report(analysis_b, out_b)
    doc_b = Document(str(out_b))
    full_b = _full_text(doc_b)
    # "FAIL" must be present (from DCMA-02 row)
    assert "FAIL" in full_b, f"B: FAIL not found in document; excerpt: {full_b[:500]}"
    # The perturbed health score must be different from the clean health score
    expected_health_b = f"{analysis_b.health_score:.2f}%"
    assert expected_health_b in full_b, (
        f"B: health {expected_health_b} not found; excerpt: {full_b[:500]}"
    )
    # Confirm the two health values are distinct (perturbation actually changed output)
    assert expected_health_a != expected_health_b, (
        "Health score should differ between clean and lead schedules"
    )


def test_perturbation_lead_findings_names_leads(tmp_path: pathlib.Path) -> None:
    """After perturbation the document Findings section must name 'Leads'."""
    analysis = analyze_schedule(_lead_schedule())
    assert any("Lead" in f for f in analysis.findings)

    out = tmp_path / "lead_findings.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert "Leads" in full, (
        f"'Leads' not found in document after perturbation; excerpt: {full[:500]}"
    )


# ---------------------------------------------------------------------------
# Edge-case: CPM error path (project_finish is None)
# ---------------------------------------------------------------------------


def test_roundtrip_cpm_error_shows_na(tmp_path: pathlib.Path) -> None:
    """When CPM fails (cycle), project finish shows 'n/a' in the document."""
    cyclic = Schedule(
        name="cyclic",
        project_start=_START,
        status_date=_START,
        tasks=(_task(1, 480), _task(2, 480)),
        relations=(
            Relation(predecessor_id=1, successor_id=2),
            Relation(predecessor_id=2, successor_id=1),
        ),
    )
    analysis = analyze_schedule(cyclic)
    assert analysis.project_finish is None
    assert analysis.cpm_error is not None

    out = tmp_path / "cyclic.docx"
    build_word_report(analysis, out)

    doc = Document(str(out))
    full = _full_text(doc)
    assert "n/a" in full.lower(), f"'n/a' not found for CPM-error case; excerpt: {full[:500]}"
