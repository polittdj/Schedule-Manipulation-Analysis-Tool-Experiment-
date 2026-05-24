"""Word report generator for Schedule Forensics (python-docx).

Public API
----------
``build_word_document(analysis) -> Document``
    Build and return the document in memory (no disk I/O). Tests use this to
    inspect the output without writing to disk.

``build_word_report(analysis, path) -> None``
    Build the document and write it to *path* as a ``.docx`` file.

Document structure
------------------
1. Title paragraph: "Schedule Forensics Report"
2. CUI banner paragraph (the ``CUI_NOTICE`` constant, bold/highlighted).
3. Executive metrics section: project finish (offset + working days), health
   score, critical path ids, driving chain ids, CPM error (if present).
4. DCMA 14-Point table: header row + one row per metric in id order.
5. Earned-Value Indices table: SPI / SPI(t) (SKIPPED rows when the schedule
   carries no earned-value data; never fabricated).
6. Findings list: one bullet per failing metric, or a "no failing metrics" note.
7. Optional Schedule Risk Analysis section (only when an
   :class:`~schedule_forensics.sra.SRAResult` is supplied): Monte-Carlo finish
   percentiles (P50/P80/P95, working days) + per-activity criticality index
   (reference-tool capability; default duration spread captioned as a tool heuristic).
8. Optional Version-to-Version Changes section (only when one or more
   :class:`~schedule_forensics.diff_engine.VersionPairDiff` are supplied): objective
   consecutive-version deltas (added/deleted, became-critical/recovered, finish/float
   shifts in working days, logic add/remove). Measured facts, not a score.
9. If any metric has ``is_extension=True``, a footnote distinguishing those rows
   as tool-original extensions (parity-honesty).

CUI / LAW 1: no schedule data leaves the machine; all writes are to the
caller-supplied local path. ``CUI_NOTICE`` is the single source of truth here.

Units: working-minute offsets are divided by 480.0 for the working-day display
value; analysis values are never recomputed -- we render exactly what
``ScheduleAnalysis`` holds (H-DRIFT-1 / LAW 2). python-docx's ``Document`` symbol
is a FACTORY; the type is ``docx.document.Document`` (imported for annotations,
constructed via ``docx.Document()``).
"""

from __future__ import annotations

import os
from typing import Any

import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from schedule_forensics.analysis import ScheduleAnalysis
from schedule_forensics.diff_engine import TaskDelta, VersionPairDiff
from schedule_forensics.metrics_common import MetricStatus
from schedule_forensics.sra import SRAResult
from schedule_forensics.trend_analysis import TrendReport

CUI_NOTICE: str = (
    "CONTROLLED UNCLASSIFIED INFORMATION (CUI) — generated locally; "
    "not for distribution. "
    "This report may contain CUI per NIST 800-171 / DFARS."
)

_MINUTES_PER_DAY: float = 480.0

_DCMA_HEADERS: list[str] = [
    "Metric ID",
    "Name",
    "Status",
    "Measured",
    "Threshold",
    "Direction",
    "Extension?",
    "Source",
    "Detail",
]

# Earned-value indices are never extensions, so no "Extension?" column.
_EV_HEADERS: list[str] = [
    "Metric ID",
    "Name",
    "Status",
    "Measured",
    "Threshold",
    "Direction",
    "Source",
    "Detail",
]


def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Set a table cell's background fill to *hex_color* (e.g. ``'C6EFCE'``)."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _status_color(status: MetricStatus) -> str:
    if status is MetricStatus.PASS:
        return "C6EFCE"  # light green
    if status is MetricStatus.FAIL:
        return "FFC7CE"  # light red
    return "FFEB9C"  # light yellow (SKIPPED)


def _add_title(doc: Document) -> None:
    title = doc.add_heading("Schedule Forensics Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cui_banner(doc: Document) -> None:
    """Add the CUI notice as a visually distinct paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(CUI_NOTICE)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(10)
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFF00")
    p_pr.append(shd)


def _add_executive_metrics(doc: Document, analysis: ScheduleAnalysis) -> None:
    doc.add_heading("Executive Metrics", level=1)

    def _row(label: str, value: object) -> None:
        para = doc.add_paragraph(style="List Bullet")
        run_label = para.add_run(f"{label}: ")
        run_label.bold = True
        para.add_run(str(value))

    if analysis.project_finish is not None:
        _row("Project Finish (working minutes)", analysis.project_finish)
        _row("Project Finish (working days)", f"{analysis.project_finish / _MINUTES_PER_DAY:.2f}")
    else:
        _row("Project Finish", "n/a (CPM could not be computed)")

    if analysis.health_score is not None:
        _row("Health Score", f"{analysis.health_score:.2f}%")
    else:
        _row("Health Score", "n/a (no runnable DCMA metrics)")

    cp_str = ", ".join(str(uid) for uid in analysis.critical_path) or "(none)"
    _row("Critical Path (UniqueIDs)", cp_str)
    dc_str = ", ".join(str(uid) for uid in analysis.driving_chain) or "(none)"
    _row("Driving Chain (UniqueIDs)", dc_str)

    if analysis.cpm_error:
        _row("CPM Error", analysis.cpm_error)


def _add_dcma_table(doc: Document, analysis: ScheduleAnalysis) -> None:
    doc.add_heading("DCMA 14-Point Assessment", level=1)

    table = doc.add_table(rows=1 + len(analysis.dcma), cols=len(_DCMA_HEADERS))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for col_idx, header in enumerate(_DCMA_HEADERS):
        cell = hdr_row.cells[col_idx]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "4472C4")

    for row_idx, metric in enumerate(analysis.dcma, start=1):
        measured_str = "" if metric.measured is None else f"{metric.measured}"
        threshold_str = "" if metric.threshold is None else f"{metric.threshold}"
        direction_str = "" if metric.direction is None else str(metric.direction)
        values = [
            metric.metric_id,
            metric.name,
            str(metric.status),
            measured_str,
            threshold_str,
            direction_str,
            "Yes" if metric.is_extension else "No",
            metric.source,
            metric.detail,
        ]
        row = table.rows[row_idx]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = val
            if col_idx == 2:
                _set_cell_bg(cell, _status_color(metric.status))


def _add_ev_table(doc: Document, analysis: ScheduleAnalysis) -> None:
    """Earned-value indices (SPI / SPI(t)); SKIPPED rows when EV data is absent."""
    doc.add_heading("Earned-Value Indices (SPI / SPI(t))", level=1)
    if not analysis.performance_indices:
        run = doc.add_paragraph().add_run("no earned-value indices")
        run.italic = True
        return

    table = doc.add_table(rows=1 + len(analysis.performance_indices), cols=len(_EV_HEADERS))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for col_idx, header in enumerate(_EV_HEADERS):
        cell = hdr_row.cells[col_idx]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "4472C4")

    for row_idx, metric in enumerate(analysis.performance_indices, start=1):
        measured_str = "" if metric.measured is None else f"{metric.measured}"
        threshold_str = "" if metric.threshold is None else f"{metric.threshold}"
        direction_str = "" if metric.direction is None else str(metric.direction)
        values = [
            metric.metric_id,
            metric.name,
            str(metric.status),
            measured_str,
            threshold_str,
            direction_str,
            metric.source,
            metric.detail,
        ]
        row = table.rows[row_idx]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = val
            if col_idx == 2:
                _set_cell_bg(cell, _status_color(metric.status))


def _add_findings(doc: Document, analysis: ScheduleAnalysis) -> None:
    doc.add_heading("Findings (Failing Metrics)", level=1)
    if analysis.findings:
        for name in analysis.findings:
            doc.add_paragraph(name, style="List Bullet")
    else:
        run = doc.add_paragraph().add_run("no failing metrics")
        run.italic = True


def _fill_header_row(table: Any, headers: list[str]) -> None:
    """Bold white-on-blue header row for *table* (row 0)."""
    hdr_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = hdr_row.cells[col_idx]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "4472C4")


def _add_trend_section(doc: Document, trends: TrendReport) -> None:
    """Multi-version trend analysis (TOOL-ORIGINAL EXTENSION): trajectory + float erosion."""
    doc.add_heading("Trend Analysis (tool-original extension)", level=1)

    if trends.finish_days_net_change is not None:
        direction = (
            "slipping"
            if trends.finish_days_net_change > 0
            else "pulling in"
            if trends.finish_days_net_change < 0
            else "no net change"
        )
        para = doc.add_paragraph()
        run = para.add_run("Finish drift: ")
        run.bold = True
        para.add_run(
            f"{trends.finish_days_first:.1f} → {trends.finish_days_last:.1f} working days "
            f"across {trends.n_versions} versions "
            f"({trends.finish_days_net_change:+.1f} days, {direction})."
        )

    doc.add_heading("Version trajectory", level=2)
    traj_headers = ["#", "Status date", "Finish (days)", "Health %", "Band", "Critical"]
    table = doc.add_table(rows=1 + len(trends.snapshots), cols=len(traj_headers))
    table.style = "Table Grid"
    _fill_header_row(table, traj_headers)
    for row_idx, snap in enumerate(trends.snapshots, start=1):
        values = [
            str(snap.index + 1),
            snap.status_date.date().isoformat() if snap.status_date else "n/a",
            f"{snap.project_finish_days:.1f}" if snap.project_finish_days is not None else "n/a",
            f"{snap.health_score:.1f}" if snap.health_score is not None else "n/a",
            snap.band,
            str(snap.n_critical),
        ]
        row = table.rows[row_idx]
        for col_idx, val in enumerate(values):
            row.cells[col_idx].text = val

    doc.add_heading("Float erosion (per task across versions)", level=2)
    tally = ", ".join(f"{band}: {count}" for band, count in trends.band_counts.items())
    doc.add_paragraph(tally)

    eroders = trends.worst_eroders(20)
    if eroders:
        er_headers = [
            "UniqueID",
            "Earliest float (d)",
            "Latest float (d)",
            "Net change (d)",
            "Trend",
        ]
        et = doc.add_table(rows=1 + len(eroders), cols=len(er_headers))
        et.style = "Table Grid"
        _fill_header_row(et, er_headers)
        for row_idx, trend in enumerate(eroders, start=1):
            values = [
                str(trend.unique_id),
                f"{trend.earliest_float_days:.1f}",
                f"{trend.latest_float_days:.1f}",
                f"{trend.net_change_days:+.1f}",
                str(trend.trend),
            ]
            row = et.rows[row_idx]
            for col_idx, val in enumerate(values):
                row.cells[col_idx].text = val

    run = doc.add_paragraph().add_run(
        "Trend Analysis is a tool-original extension (the cross-version trajectory framing and "
        "float-erosion bands); per-version finish and health values are objective, but the trend "
        "layer must not be presented as reference-tool parity (parity-honesty rule -- CLAUDE.md)."
    )
    run.italic = True


def _add_sra_section(doc: Document, sra: SRAResult) -> None:
    """Monte-Carlo Schedule Risk Analysis: finish percentiles + criticality index.

    The SRA method (finish distribution + criticality index) is a reference-tool
    capability (Acumen Fuse Risk tab / Primavera Risk Analysis) -- NOT flagged a
    tool-original extension. The DEFAULT duration spread that seeds it IS the
    tool's own heuristic (source-pending) and is captioned as such. Finish offsets
    render in working days (offset / 480.0), consistent with the rest of the report."""
    doc.add_heading("Schedule Risk Analysis (Monte Carlo)", level=1)

    def _row(label: str, value: object) -> None:
        para = doc.add_paragraph(style="List Bullet")
        run_label = para.add_run(f"{label}: ")
        run_label.bold = True
        para.add_run(str(value))

    _row("Monte-Carlo iterations", sra.iterations)
    _row(
        "Deterministic finish (working days)",
        f"{sra.deterministic_finish / _MINUTES_PER_DAY:.1f}",
    )
    _row("Mean finish (working days)", f"{sra.mean_finish / _MINUTES_PER_DAY:.1f}")

    doc.add_heading("Finish percentiles (working days)", level=2)
    pct_headers = ["Percentile", "Finish (working days)"]
    table = doc.add_table(rows=1 + 3, cols=len(pct_headers))
    table.style = "Table Grid"
    _fill_header_row(table, pct_headers)
    for row_idx, (label, offset) in enumerate(
        (("P50", sra.p50), ("P80", sra.p80), ("P95", sra.p95)), start=1
    ):
        row = table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = f"{offset / _MINUTES_PER_DAY:.1f}"

    doc.add_heading("Criticality index (per activity)", level=2)
    ranked = sorted(sra.criticality_index.items(), key=lambda kv: (-kv[1], kv[0]))
    on_path = [(uid, ci) for uid, ci in ranked if ci > 0.0]
    if on_path:
        ci_headers = ["UniqueID", "Criticality index (%)"]
        ct = doc.add_table(rows=1 + len(on_path[:25]), cols=len(ci_headers))
        ct.style = "Table Grid"
        _fill_header_row(ct, ci_headers)
        for row_idx, (uid, ci) in enumerate(on_path[:25], start=1):
            row = ct.rows[row_idx]
            row.cells[0].text = str(uid)
            row.cells[1].text = f"{100.0 * ci:.1f}"
    else:
        run = doc.add_paragraph().add_run("no activity reached the critical path in any iteration")
        run.italic = True

    run = doc.add_paragraph().add_run(
        "The SRA method (finish distribution + criticality index) mirrors Acumen Fuse / "
        "Primavera Risk Analysis (reference parity). The default duration spread that seeds it is "
        "the tool's own heuristic (source-pending); in an engagement the analyst supplies "
        "per-activity risk ranges. The distribution models duration uncertainty only."
    )
    run.italic = True


def _notable_deltas(pair: VersionPairDiff) -> list[TaskDelta]:
    """Task deltas where anything actually changed, most-moved (by finish shift) first."""
    changed = [
        d
        for d in pair.task_deltas
        if d.became_critical
        or d.recovered
        or d.finish_shift_minutes
        or d.total_float_delta_minutes
        or d.predecessors_added
        or d.predecessors_removed
    ]
    changed.sort(key=lambda d: (-abs(d.finish_shift_minutes), d.unique_id))
    return changed


def _add_diff_section(doc: Document, diff: tuple[VersionPairDiff, ...]) -> None:
    """Objective version-to-version deltas (NOT an extension -- comparative facts).

    Per consecutive pair: the objective counts and the most-moved task deltas, with
    finish/float shifts in working days. These are measured facts (the Acumen
    ProjectTimeNow comparative pattern), so no threshold is applied."""
    doc.add_heading("Version-to-Version Changes", level=1)
    for pair in diff:
        doc.add_heading(
            f"{pair.previous_status.date().isoformat()} → {pair.current_status.date().isoformat()}",
            level=2,
        )
        n_bc = sum(1 for d in pair.task_deltas if d.became_critical)
        n_rec = sum(1 for d in pair.task_deltas if d.recovered)
        n_logic = sum(
            len(d.predecessors_added) + len(d.predecessors_removed) for d in pair.task_deltas
        )
        para = doc.add_paragraph()
        para.add_run(
            f"Tasks added/deleted: {len(pair.added_ids)}/{len(pair.deleted_ids)}; "
            f"became critical/recovered: {n_bc}/{n_rec}; logic links added+removed: {n_logic}."
        )

        notable = _notable_deltas(pair)
        if not notable:
            run = doc.add_paragraph().add_run("no task-level changes between these versions")
            run.italic = True
            continue
        headers = ["UniqueID", "Finish shift (d)", "Float delta (d)", "Flag", "Preds +", "Preds -"]
        table = doc.add_table(rows=1 + len(notable), cols=len(headers))
        table.style = "Table Grid"
        _fill_header_row(table, headers)
        for row_idx, d in enumerate(notable, start=1):
            flag = "became critical" if d.became_critical else "recovered" if d.recovered else ""
            values = [
                str(d.unique_id),
                f"{d.finish_shift_minutes / _MINUTES_PER_DAY:+.1f}",
                f"{d.total_float_delta_minutes / _MINUTES_PER_DAY:+.1f}",
                flag,
                ", ".join(str(p) for p in d.predecessors_added),
                ", ".join(str(p) for p in d.predecessors_removed),
            ]
            row = table.rows[row_idx]
            for col_idx, val in enumerate(values):
                row.cells[col_idx].text = val

    run = doc.add_paragraph().add_run(
        "Objective version deltas (the Acumen ProjectTimeNow / ProjectPreviousTimeNow comparative "
        "pattern): measured facts, not a score; no threshold is applied. Positive finish shift = "
        "task moved later; positive float delta = float gained."
    )
    run.italic = True


def _add_extension_footnote(doc: Document, analysis: ScheduleAnalysis) -> None:
    extension_ids = [m.metric_id for m in analysis.dcma if m.is_extension]
    if not extension_ids:
        return
    doc.add_heading("Tool-Original Extensions", level=2)
    run = doc.add_paragraph().add_run(
        f"The following metrics ({', '.join(extension_ids)}) are tool-original extensions: "
        "capabilities beyond what Acumen Fuse, Steelray/SSI, or Microsoft Project produce "
        "natively. They are labeled 'Yes' in the Extension? column and must not be presented "
        "as reference-tool parity (parity-honesty rule -- CLAUDE.md)."
    )
    run.italic = True


def build_word_document(
    analysis: ScheduleAnalysis,
    trends: TrendReport | None = None,
    sra: SRAResult | None = None,
    diff: tuple[VersionPairDiff, ...] = (),
) -> Document:
    """Build and return a python-docx ``Document`` from *analysis* (no disk I/O).

    When *trends* is given and spans 2+ versions, a **Trend Analysis** section
    (tool-original extension: version trajectory + float erosion) is added. When
    *sra* is given, a **Schedule Risk Analysis** section (Monte-Carlo finish
    percentiles + criticality index; reference-tool capability) is added. When
    *diff* has any version pairs, a **Version-to-Version Changes** section
    (objective consecutive-version deltas) is added. All precede the extensions
    footnote."""
    doc = docx.Document()
    _add_title(doc)
    _add_cui_banner(doc)
    _add_executive_metrics(doc, analysis)
    _add_dcma_table(doc, analysis)
    _add_ev_table(doc, analysis)
    _add_findings(doc, analysis)
    if trends is not None and trends.n_versions > 1:
        _add_trend_section(doc, trends)
    if sra is not None:
        _add_sra_section(doc, sra)
    if diff:
        _add_diff_section(doc, diff)
    _add_extension_footnote(doc, analysis)
    return doc


def build_word_report(
    analysis: ScheduleAnalysis,
    path: str | os.PathLike[str],
    trends: TrendReport | None = None,
    sra: SRAResult | None = None,
    diff: tuple[VersionPairDiff, ...] = (),
) -> None:
    """Write a Word ``.docx`` report for *analysis* to *path* (all data stays local)."""
    build_word_document(analysis, trends=trends, sra=sra, diff=diff).save(os.fspath(path))
